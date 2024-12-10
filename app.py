from io import BytesIO
import streamlit as st
import pandas as pd
import altair as alt
from sidebar import create_sidebar
from bubbles import render_bubbles
from theme import set_page_mode
from header import render_header
from data_utils import clean_data
from aitest import generate_comments
from PIL import Image
from docx import Document
from docx.shared import Pt, Inches
import tempfile
import altair_saver as alt_save
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Main app
st.set_page_config(layout="wide")
# Add the custom CSS and HTML to Streamlit
render_bubbles()

if "text_area_content" not in st.session_state: 
    st.session_state.text_area_content = ""

if "chart_data" not in st.session_state:
    st.session_state.chart_data = None, None, None, None, None, None

if "agg_method" not in st.session_state:
    st.session_state.agg_method = None

uploaded_file = None  # Initialize file

# Set page mode
set_page_mode("Dark")

# Render header
render_header()
uploaded_file = st.file_uploader("Upload your CSV file", type=["csv"])

if uploaded_file:
    try:
        # Load uploaded CSV
        df = pd.read_csv(uploaded_file)
        if "df" not in st.session_state:
            st.session_state.df = df
        
        currency_columns = []
        for column in df.columns: 
            if isinstance(df[column].iloc[0], str) and df[column].iloc[0].startswith('$'): 
                currency_columns.append(column)
        
        for column in currency_columns: 
            df[column] = df[column].replace({'\$': '', ',': ''}, regex=True).astype(float)

        clean_data_options = create_sidebar(uploaded_file, st.session_state.df)

        # Number of rows to display
        num_rows = st.slider("Number of rows to display in preview", min_value=5, max_value=100, value=10, step=5)

        # Apply cleaning options only after file upload
        if clean_data_options:
            df = clean_data(df, clean_data_options)
            st.session_state.df = df

        # Display cleaned data preview
        st.write(f"### Cleaned Data Preview ({num_rows} rows)")
        st.write(df.head(num_rows))

        # Chart Rendering Logic
        chart_type, x_axis, y_axis, x_axis_dtype, y_axis_dtype, df_counts = st.session_state.chart_data
        agg_method = st.session_state.agg_method

        def save_chart(chart): 
            buffer = BytesIO() 
            alt_save.save(chart, buffer, fmt='png') 
            buffer.seek(0) 
            return Image.open(buffer)
        
        
        if chart_type:
            y_axis_name = 'counts' if 'counts' in df_counts else y_axis
            if chart_type == "Bar Chart":
                chart = alt.Chart(df_counts).mark_bar().encode(
                    x=alt.X(f'{x_axis}:N', title=x_axis),
                    y=alt.Y(f'{y_axis_name}:Q', title=y_axis_name),
                    tooltip=[x_axis, y_axis_name]
                ).properties(
                    title=f"Bar Chart of {agg_method} {y_axis} vs {x_axis}",
                    width=800,
                    height=400
                ).interactive()

                st.altair_chart(chart, use_container_width=True)
            elif chart_type == "Line Chart":
                chart = alt.Chart(df_counts).mark_line().encode(
                    x=alt.X(f'{x_axis}:N', title=x_axis),
                    y=alt.Y(f'{y_axis_name}:Q', title=y_axis_name),
                    tooltip=[x_axis, y_axis_name]
                ).properties(
                    title=f"Line Chart of {agg_method} {y_axis} over {x_axis}",
                    width=800,
                    height=400
                ).interactive()

                st.altair_chart(chart, use_container_width=True)
            elif chart_type == "Scatter Plot":
                chart = alt.Chart(df_counts).mark_circle().encode(
                    x=alt.X(f'{x_axis}:N', title=x_axis),
                    y=alt.Y(f'{y_axis_name}:Q', title=y_axis_name),
                    tooltip=[x_axis, y_axis_name]
                ).properties(
                    title=f"Scatter Plot of {agg_method} {y_axis} vs {x_axis}",
                    width=800,
                    height=400
                ).interactive()

                st.altair_chart(chart, use_container_width=True)
            elif chart_type == "Table":
                st.write("### Full Data Table")
                st.write(df)
            elif chart_type == "Grouped Bar Chart": 
                st.write("### Grouped Bar Chart") 
                chart = alt.Chart(df_counts).mark_bar().encode( 
                    x=alt.X(f'{x_axis}:N', title=x_axis),
                    y=alt.Y('sum(counts):Q', title='Counts'), 
                    color=f'{y_axis}:N', 
                    # column=f'{y_axis}:N', 
                    # tooltip=[x_axis, y_axis, 'counts'] 
                    ).properties( 
                        title=f"Grouped Bar Chart of {agg_method} {y_axis} vs {x_axis}", 
                        width=100, 
                        height=400 
                    ).configure_axis( 
                        labelFontSize=12, 
                        titleFontSize=14 
                        ).interactive() 
                
                st.altair_chart(chart, use_container_width=True)

            elif chart_type == "Layered Histogram": 
                st.write("### Layered Histogram") 
                chart = alt.Chart(df_counts).mark_bar(opacity=0.5).encode( 
                    x=alt.X(f'{x_axis}:Q', bin=alt.Bin(maxbins=30), title=x_axis), 
                    y=alt.Y('count()', title='Count'), 
                    color=f'{y_axis}:N', 
                    tooltip=[x_axis, y_axis] 
                    ).properties( 
                        title=f"Layered Histogram of {y_axis} vs {x_axis}", 
                        width=800, height=400 
                        ).interactive() 
                st.altair_chart(chart, use_container_width=True)

            # Add comments and download options
            st.write("### Comments or Observations")

            comments = st.text_area(
                "Add your comments or observations about the visualization here:",
                placeholder="Enter your thoughts...",
                height=200,
                value=st.session_state.text_area_content
            )
            if st.button("Generate with AI 🤖") and chart is not None:
                with st.spinner('Generating comments with AI...'):
                    chart_json = chart.to_json()
                    # image = save_chart(chart)
                    response = generate_comments(chart_json)
                    st.session_state.text_area_content = f"{response}"
                    st.rerun()

            st.write("### Generate Report")
            if st.button("Generate Report"):
                if comments.strip():
                    with st.spinner('Generating the report...'):
                        # Create a Word document
                        doc = Document()
                        doc.add_heading('Report from PogiOnly', 0)
                        
                        # Add image
                        buffer = BytesIO()
                        chart.save(buffer, format="png")
                        buffer.seek(0)
                        image = Image.open(buffer)
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_file:
                            image.save(tmp_file.name)
                            doc.add_heading('Generated Image:', level=1)
                            p = doc.add_paragraph()
                            run = p.add_run()
                            run.add_picture(tmp_file.name, width=Inches(4))  # Adjust the width as needed
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        
                        # Add comments to the document
                        doc.add_heading('User Comments:', level=1)
                        for line in comments.split('\n'):
                            paragraph = doc.add_paragraph()
                            run = paragraph.add_run(line)
                            run.font.size = Pt(12)

                        # Save the document to a BytesIO object
                        buffer = BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)

                    st.success("Report Generated!")
                    st.download_button(
                        label="Download DOCX Report",
                        data=buffer,
                        file_name="report_with_comments.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.warning("Please add comments before generating the report.")

    except Exception as e:
        st.error(f"An error occurred while processing the file: {e}")
else:
    st.warning("🔴 Please upload a CSV file to start.")
    pass
